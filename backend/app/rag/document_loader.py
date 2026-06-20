from pathlib import Path
from typing import List

from langchain_core.documents import Document
from langchain_community.document_loaders import PyPDFLoader
from docx import Document as DocxDocument


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


def clean_text(text: str) -> str:
    """
    Basic text cleanup.
    Keeps the content readable without over-processing it.
    """
    if not text:
        return ""

    lines = text.splitlines()
    cleaned_lines = []

    for line in lines:
        stripped_line = line.strip()
        if stripped_line:
            cleaned_lines.append(stripped_line)

    return "\n".join(cleaned_lines)


def load_pdf(file_path: Path) -> List[Document]:
    """
    Loads PDF file using LangChain PyPDFLoader.
    Usually returns one Document per page.
    """
    loader = PyPDFLoader(str(file_path))
    documents = loader.load()

    cleaned_documents = []

    for doc in documents:
        page_content = clean_text(doc.page_content)

        if not page_content:
            continue

        metadata = dict(doc.metadata)
        metadata["source"] = str(file_path)
        metadata["file_name"] = file_path.name
        metadata["file_type"] = ".pdf"

        cleaned_documents.append(
            Document(
                page_content=page_content,
                metadata=metadata
            )
        )

    return cleaned_documents


def load_docx(file_path: Path) -> List[Document]:
    """
    Loads DOCX file using python-docx.
    Extracts normal paragraphs and basic table text.
    """
    docx_file = DocxDocument(str(file_path))

    text_parts = []

    for paragraph in docx_file.paragraphs:
        if paragraph.text and paragraph.text.strip():
            text_parts.append(paragraph.text.strip())

    for table in docx_file.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = clean_text(cell.text)
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                text_parts.append(" | ".join(row_text))

    full_text = clean_text("\n".join(text_parts))

    if not full_text:
        return []

    return [
        Document(
            page_content=full_text,
            metadata={
                "source": str(file_path),
                "file_name": file_path.name,
                "file_type": ".docx"
            }
        )
    ]


def load_text_file(file_path: Path) -> List[Document]:
    """
    Loads TXT or Markdown files.
    """
    text = file_path.read_text(encoding="utf-8", errors="ignore")
    cleaned = clean_text(text)

    if not cleaned:
        return []

    return [
        Document(
            page_content=cleaned,
            metadata={
                "source": str(file_path),
                "file_name": file_path.name,
                "file_type": file_path.suffix.lower()
            }
        )
    ]


def load_document(file_path: str) -> List[Document]:
    """
    Main loader function.
    Detects file type and sends file to the correct loader.
    """

    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"File does not exist: {file_path}")

    extension = path.suffix.lower()

    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file extension '{extension}'. "
            f"Supported extensions are: {', '.join(sorted(SUPPORTED_EXTENSIONS))}"
        )

    if extension == ".pdf":
        return load_pdf(path)

    if extension == ".docx":
        return load_docx(path)

    if extension in {".txt", ".md"}:
        return load_text_file(path)

    return []